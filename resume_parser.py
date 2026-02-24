"""
Resume Parser Module — Multi-Stage Fault-Tolerant Pipeline Orchestrator

Detects file type via magic bytes, performs layered text extraction with
layout preservation (bounding boxes, font attributes, tables, columns),
triggers OCR when native extraction is incomplete/corrupted, and returns
a deterministic ParseResult with full traceability.
"""

import io
import os
import re
import sys
import tempfile
import logging
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Tuple

logger = logging.getLogger(__name__)


# ── Data Classes ──────────────────────────────────────────

@dataclass
class WordInfo:
    """A single word with bounding box and font metadata."""
    text: str
    x0: float = 0.0
    y0: float = 0.0
    x1: float = 0.0
    y1: float = 0.0
    font_name: str = ""
    font_size: float = 0.0
    page: int = 0


@dataclass
class PageData:
    """Per-page extraction data with layout metadata."""
    page_number: int
    text: str = ""
    words: List[WordInfo] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)
    width: float = 0.0
    height: float = 0.0


@dataclass
class ParseResult:
    """Complete result from the parsing pipeline."""
    raw_text: str = ""
    cleaned_text: str = ""
    pages: List[PageData] = field(default_factory=list)
    extraction_method: str = "unknown"  # native, ocr, hybrid
    file_type: str = "unknown"          # pdf, doc, docx, txt
    anomalies: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


# ── File Type Detection ──────────────────────────────────

def detect_file_type(file_bytes: bytes, filename: str = "") -> str:
    """
    Detect file type from magic bytes, falling back to extension.
    Deterministic — no randomness.
    """
    if len(file_bytes) < 4:
        return "txt"

    # PDF: starts with %PDF
    if file_bytes[:4] == b"%PDF":
        return "pdf"

    # DOC (OLE2 Compound Document): starts with D0 CF 11 E0
    if file_bytes[:4] == b"\xd0\xcf\x11\xe0":
        return "doc"

    # DOCX (ZIP/PK archive containing word/document.xml)
    if file_bytes[:2] == b"PK":
        return "docx"

    # Fallback to extension
    ext = os.path.splitext(filename)[1].lower() if filename else ""
    if ext in (".pdf", ".doc", ".docx", ".txt"):
        return ext[1:]

    return "txt"


# ── Text Corruption Heuristics ───────────────────────────

def _assess_text_quality(text: str, file_size: int) -> Tuple[float, List[str]]:
    """
    Assess extracted text quality. Returns (quality_score 0-1, anomalies).
    Deterministic thresholds.
    """
    anomalies = []

    if not text or not text.strip():
        anomalies.append("CORRUPTION: Extraction yielded empty text")
        return 0.0, anomalies

    total_chars = len(text)
    stripped = text.strip()

    # 1. Character length vs file size
    if file_size > 0:
        char_per_byte = total_chars / file_size
        if char_per_byte < 0.002:  # <1 word per 500 bytes
            anomalies.append(f"CORRUPTION: Very low text density ({char_per_byte:.4f} chars/byte)")

    # 2. Unknown/replacement character ratio
    unknown_chars = sum(1 for c in stripped if c in ('\ufffd', '\x00', '\x01', '\x02'))
    if total_chars > 0:
        unknown_ratio = unknown_chars / total_chars
        if unknown_ratio > 0.15:
            anomalies.append(f"CORRUPTION: High unknown char ratio ({unknown_ratio:.1%})")

    # 3. Word count sanity
    words = stripped.split()
    if len(words) < 10 and file_size > 5000:
        anomalies.append(f"CORRUPTION: Only {len(words)} words from {file_size} byte file")

    # 4. Encoding anomalies — sequences like Ã©, â€™ suggest double-encoding
    mojibake_patterns = [r'Ã[\x80-\xbf]', r'â€[™""]', r'Ã¢â‚¬']
    for pat in mojibake_patterns:
        if re.search(pat, stripped):
            anomalies.append("CORRUPTION: Mojibake encoding artifacts detected")
            break

    # Score: 1.0 = perfect, 0.0 = completely corrupt
    score = 1.0
    if len(words) < 10:
        score -= 0.5
    if total_chars < 50:
        score -= 0.3
    for a in anomalies:
        if "unknown char ratio" in a:
            score -= 0.4
        elif "Mojibake" in a:
            score -= 0.2
        elif "text density" in a:
            score -= 0.3

    return max(0.0, min(1.0, score)), anomalies


# ── PDF Extraction (pdfplumber + bounding boxes) ─────────

def _extract_pdf_pdfplumber(file_bytes: bytes) -> Tuple[str, List[PageData], List[str]]:
    """
    Primary PDF extraction using pdfplumber with full layout metadata.
    Returns (text, pages, anomalies).
    """
    import pdfplumber

    pages_data = []
    text_parts = []
    anomalies = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                pd = PageData(
                    page_number=page_num + 1,
                    width=float(page.width),
                    height=float(page.height)
                )

                # Extract words with bounding boxes and font info
                try:
                    words_raw = page.extract_words(
                        x_tolerance=2, y_tolerance=3,
                        extra_attrs=["fontname", "size"]
                    )
                    for w in words_raw:
                        pd.words.append(WordInfo(
                            text=w.get("text", ""),
                            x0=float(w.get("x0", 0)),
                            y0=float(w.get("top", 0)),
                            x1=float(w.get("x1", 0)),
                            y1=float(w.get("bottom", 0)),
                            font_name=w.get("fontname", ""),
                            font_size=float(w.get("size", 0)),
                            page=page_num + 1
                        ))
                except Exception as e:
                    anomalies.append(f"PAGE_{page_num+1}: Word extraction failed: {e}")

                # Extract text with layout preservation
                try:
                    page_text = page.extract_text(
                        layout=True, x_tolerance=2, y_tolerance=3
                    )
                    if page_text:
                        pd.text = page_text
                        text_parts.append(page_text)
                except Exception as e:
                    anomalies.append(f"PAGE_{page_num+1}: Text extraction failed: {e}")

                # Extract tables
                try:
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            cleaned_table = []
                            for row in table:
                                cleaned_row = [
                                    (cell or "").strip() for cell in row
                                ]
                                cleaned_table.append(cleaned_row)
                            pd.tables.append(cleaned_table)
                except Exception as e:
                    anomalies.append(f"PAGE_{page_num+1}: Table extraction failed: {e}")

                pages_data.append(pd)

    except Exception as e:
        anomalies.append(f"PDF_OPEN: pdfplumber failed: {e}")

    full_text = "\n".join(text_parts).strip()
    return full_text, pages_data, anomalies


def _extract_pdf_fitz(file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Secondary PDF extraction using PyMuPDF (fitz) for fast text.
    Used for multi-engine confidence comparison.
    Returns (text, anomalies).
    """
    anomalies = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text = page.get_text("text")
            if text:
                text_parts.append(text)
        doc.close()
        return "\n".join(text_parts).strip(), anomalies
    except ImportError:
        anomalies.append("ENGINE: PyMuPDF not available for secondary extraction")
        return "", anomalies
    except Exception as e:
        anomalies.append(f"ENGINE: PyMuPDF extraction failed: {e}")
        return "", anomalies


def _merge_extraction_results(
    primary_text: str, primary_pages: List[PageData],
    secondary_text: str, file_size: int
) -> Tuple[str, str, List[str]]:
    """
    Confidence-weighted merge of primary and secondary extraction results.
    Returns (best_text, method_used, anomalies).
    """
    anomalies = []

    primary_score, primary_anomalies = _assess_text_quality(primary_text, file_size)
    secondary_score, secondary_anomalies = _assess_text_quality(secondary_text, file_size)

    anomalies.extend(primary_anomalies)

    if primary_score >= 0.7:
        return primary_text, "pdfplumber", anomalies

    if secondary_score > primary_score + 0.15:
        anomalies.append(
            f"ENGINE_SWITCH: Using PyMuPDF (score {secondary_score:.2f}) over "
            f"pdfplumber (score {primary_score:.2f})"
        )
        return secondary_text, "pymupdf", anomalies

    # If both are poor, concat unique content
    if primary_score < 0.3 and secondary_score < 0.3:
        anomalies.append("ENGINE: Both engines produced low-quality text")
        # Use whichever has more content
        if len(secondary_text) > len(primary_text) * 1.5:
            return secondary_text, "pymupdf", anomalies

    return primary_text, "pdfplumber", anomalies


# ── PDF Extraction Orchestrator ──────────────────────────

def parse_pdf(file_bytes: bytes) -> Tuple[str, List[PageData], str, List[str]]:
    """
    Multi-engine PDF extraction with confidence-based merging.
    Returns (text, pages, method, anomalies).
    """
    anomalies = []

    # Primary engine: pdfplumber (with bounding boxes)
    primary_text, pages, plumber_anomalies = _extract_pdf_pdfplumber(file_bytes)
    anomalies.extend(plumber_anomalies)

    # Secondary engine: PyMuPDF (fast text)
    secondary_text, fitz_anomalies = _extract_pdf_fitz(file_bytes)
    anomalies.extend(fitz_anomalies)

    # Confidence-weighted merge
    best_text, method, merge_anomalies = _merge_extraction_results(
        primary_text, pages, secondary_text, len(file_bytes)
    )
    anomalies.extend(merge_anomalies)

    return best_text, pages, method, anomalies


# ── DOCX Extraction ──────────────────────────────────────

def _iter_block_items(parent):
    """Yield each paragraph and table child within *parent*, in document order."""
    from docx.document import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent must be Document or _Cell")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_docx(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Extract text from DOCX, preserving document order. Returns (text, anomalies)."""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    anomalies = []
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        anomalies.append(f"DOCX_OPEN: Failed to open: {e}")
        return "", anomalies

    text_parts = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                text_parts.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

    return "\n".join(text_parts), anomalies


# ── DOC Extraction ───────────────────────────────────────

def parse_doc(file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Extract text from legacy .doc files.
    Strategy 1: COM automation on Windows (requires Word)
    Strategy 2: antiword CLI fallback
    Returns (text, anomalies).
    """
    anomalies = []

    # Strategy 1: Windows COM automation
    if sys.platform == "win32":
        try:
            import win32com.client
            import pythoncom

            pythoncom.CoInitialize()
            try:
                # Write bytes to temp .doc file
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_doc_path = tmp.name

                tmp_docx_path = tmp_doc_path + "x"

                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False

                try:
                    doc = word.Documents.Open(tmp_doc_path)
                    doc.SaveAs2(tmp_docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault (docx)
                    doc.Close()
                finally:
                    word.Quit()

                # Now parse the converted DOCX
                with open(tmp_docx_path, "rb") as f:
                    docx_bytes = f.read()

                text, docx_anomalies = parse_docx(docx_bytes)
                anomalies.extend(docx_anomalies)

                # Cleanup temp files
                for p in (tmp_doc_path, tmp_docx_path):
                    try:
                        os.unlink(p)
                    except OSError:
                        pass

                return text, anomalies
            finally:
                pythoncom.CoUninitialize()

        except ImportError:
            anomalies.append("DOC: pywin32 not installed, trying antiword fallback")
        except Exception as e:
            anomalies.append(f"DOC: COM automation failed: {e}")

    # Strategy 2: antiword CLI
    try:
        import subprocess

        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            result = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip(), anomalies
            else:
                anomalies.append(f"DOC: antiword returned code {result.returncode}")
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    except FileNotFoundError:
        anomalies.append("DOC: antiword not found on system PATH")
    except Exception as e:
        anomalies.append(f"DOC: antiword failed: {e}")

    # Strategy 3: Last resort — try reading as text
    anomalies.append("DOC: All extraction methods failed, attempting raw text decode")
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        # Filter out control characters common in binary DOC format
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        return text.strip(), anomalies
    except Exception:
        return "", anomalies


# ── Plain Text Extraction ────────────────────────────────

def parse_txt(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Read plain text with encoding detection. Returns (text, anomalies)."""
    anomalies = []

    # Try UTF-8 first, then common encodings
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding)
            if encoding != "utf-8":
                anomalies.append(f"ENCODING: Decoded as {encoding} (not UTF-8)")
            return text.strip(), anomalies
        except (UnicodeDecodeError, UnicodeError):
            continue

    anomalies.append("ENCODING: All standard decodings failed, using lossy UTF-8")
    return file_bytes.decode("utf-8", errors="replace").strip(), anomalies


# ── Main Pipeline Orchestrator ───────────────────────────

def parse_resume(file, filename: str = None) -> ParseResult:
    """
    Multi-stage, fault-tolerant resume parsing pipeline.

    Parameters
    ----------
    file : str, bytes, or file-like
        Path, raw bytes, or uploaded file object.
    filename : str, optional
        Original filename for extension-based fallback detection.

    Returns
    -------
    ParseResult
        Structured result with raw/cleaned text, pages, bounding boxes,
        extraction method, anomalies, and metadata.
    """
    result = ParseResult()

    # ── Step 1: Read file to bytes ──
    if isinstance(file, str):
        filename = filename or file
        try:
            with open(file, "rb") as f:
                file_bytes = f.read()
        except Exception as e:
            result.anomalies.append(f"FILE_READ: Cannot read '{file}': {e}")
            return result
    elif isinstance(file, bytes):
        file_bytes = file
    elif hasattr(file, "read"):
        file_bytes = file.read()
        if isinstance(file_bytes, str):
            file_bytes = file_bytes.encode("utf-8")
        filename = filename or getattr(file, "name", "")
    else:
        result.anomalies.append("FILE_READ: Unsupported file type")
        return result

    if not file_bytes:
        result.anomalies.append("FILE_READ: Empty file")
        return result

    # ── Step 2: Detect file type via magic bytes ──
    result.file_type = detect_file_type(file_bytes, filename or "")
    logger.info(f"Detected file type: {result.file_type} ({len(file_bytes)} bytes)")

    # ── Step 3: Extract text based on file type ──
    raw_text = ""
    pages = []
    method = "native"

    if result.file_type == "pdf":
        raw_text, pages, method, anomalies = parse_pdf(file_bytes)
        result.anomalies.extend(anomalies)
        result.pages = pages

        # ── Step 3b: OCR fallback for PDFs ──
        quality_score, quality_anomalies = _assess_text_quality(
            raw_text, len(file_bytes)
        )
        result.anomalies.extend(quality_anomalies)

        if quality_score < 0.5:
            logger.warning(
                f"PDF text quality low ({quality_score:.2f}). Triggering OCR fallback..."
            )
            try:
                from ocr_engine import extract_text_from_pdf_bytes, OcrResult

                ocr_result = extract_text_from_pdf_bytes(file_bytes)

                if isinstance(ocr_result, OcrResult):
                    ocr_text = ocr_result.text
                    result.anomalies.extend(ocr_result.anomalies)
                    result.metadata["ocr_confidence"] = ocr_result.mean_confidence
                else:
                    ocr_text = ocr_result if isinstance(ocr_result, str) else ""

                if ocr_text and len(ocr_text) > len(raw_text):
                    if quality_score < 0.3:
                        raw_text = ocr_text
                        method = "ocr"
                    else:
                        # Hybrid: use native bounding boxes + OCR text
                        raw_text = ocr_text
                        method = "hybrid"
                    logger.info(f"OCR fallback successful ({len(ocr_text)} chars)")
                else:
                    result.anomalies.append("OCR: OCR did not improve extraction")

            except ImportError:
                result.anomalies.append("OCR: ocr_engine module not available")
            except Exception as e:
                result.anomalies.append(f"OCR: Fallback failed: {e}")

    elif result.file_type == "docx":
        raw_text, anomalies = parse_docx(file_bytes)
        result.anomalies.extend(anomalies)

    elif result.file_type == "doc":
        raw_text, anomalies = parse_doc(file_bytes)
        result.anomalies.extend(anomalies)

    elif result.file_type == "txt":
        raw_text, anomalies = parse_txt(file_bytes)
        result.anomalies.extend(anomalies)

    else:
        raw_text, anomalies = parse_txt(file_bytes)
        result.anomalies.extend(anomalies)

    result.raw_text = raw_text
    result.extraction_method = method

    # ── Step 4: Layout Repair & Normalization ──
    cleaned_text = raw_text
    try:
        from layout_processor import repair_layout

        repair_result = repair_layout(raw_text, pages)

        if isinstance(repair_result, tuple):
            cleaned_text, repair_anomalies = repair_result
            result.anomalies.extend(repair_anomalies)
        else:
            cleaned_text = repair_result

    except ImportError:
        result.anomalies.append("LAYOUT: layout_processor module not available")
    except Exception as e:
        result.anomalies.append(f"LAYOUT: Repair failed: {e}")

    result.cleaned_text = cleaned_text

    # ── Step 5: Finalize metadata ──
    result.metadata.update({
        "file_size_bytes": len(file_bytes),
        "char_count_raw": len(raw_text),
        "char_count_cleaned": len(cleaned_text),
        "page_count": len(pages) if pages else 0,
        "table_count": sum(len(p.tables) for p in pages) if pages else 0,
        "word_count": len(cleaned_text.split()) if cleaned_text else 0,
    })

    logger.info(
        f"Parse complete: {result.extraction_method} | "
        f"{result.metadata.get('char_count_cleaned', 0)} chars | "
        f"{len(result.anomalies)} anomalies"
    )

    return result


# ── Backward Compatibility Wrapper ───────────────────────

def parse_resume_text(file, filename: str = None) -> str:
    """
    Backward-compatible wrapper that returns just the cleaned text string.
    Existing code calling the old parse_resume() can use this.
    """
    result = parse_resume(file, filename)
    return result.cleaned_text or result.raw_text
